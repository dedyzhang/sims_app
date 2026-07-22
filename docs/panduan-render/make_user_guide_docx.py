"""Buat Word user guide: Arena Belajar + Asisten Guru (update)."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "PANDUAN-ARENA-BELAJAR-DAN-ASISTEN-GURU.docx"
IMG = ROOT / "public" / "images" / "panduan"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)


def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def p(text, bold=False, size=11):
    para = doc.add_paragraph()
    r = para.add_run(text)
    font(r, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(6)
    return para


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    r = para.add_run(text)
    font(r)
    para.paragraph_format.space_after = Pt(2)


def num(text):
    para = doc.add_paragraph(style="List Number")
    r = para.add_run(text)
    font(r)
    para.paragraph_format.space_after = Pt(2)


def add_img(name, width_cm=14):
    path = IMG / name
    if path.exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph()
        r = cap.add_run(f"Gambar: {name}")
        font(r, size=9, color=(100, 100, 100))
        cap.paragraph_format.space_after = Pt(10)


title = doc.add_paragraph()
r = title.add_run("PANDUAN PENGGUNA — ARENA BELAJAR & ASISTEN GURU")
font(r, size=18, bold=True, color=(15, 76, 129))

sub = doc.add_paragraph()
r = sub.add_run("Update Juli 2026 · SIMS · Dilengkapi video storyboard di aplikasi")
font(r, size=11, color=(80, 80, 80))

meta = doc.add_paragraph()
r = meta.add_run(
    "Di aplikasi: Bantuan → Panduan Visual. "
    "Video: /videos/panduan/arenabelajar.mp4, arenakuis.mp4, arenamisi.mp4, ai.mp4"
)
font(r, size=10, color=(100, 100, 100))

h("1. Ringkas alur yang disarankan", 1)
p("Asisten Guru (buat soal) → Kirim ke Arena → periksa kunci → Terbitkan → siswa main (solo/live/template) → transfer nilai.")
bullet("Asisten Guru: API key Gemini pribadi + folder Nalar Guru & Kuota.")
bullet("Arena: kuis + misi; solo acak; skin Ular tangga; salin ke kelas lain.")

h("2. Arena Belajar", 1)
p("Lokasi: Akademik → Ruang Kelas → pilih ruang mapel → tab Arena Belajar.", bold=True)
add_img("arena-belajar.png")

h("2.1 Siapa yang memakai", 2)
bullet("Guru: buat/terbitkan, template, salin, live, misi, hasil, transfer nilai.")
bullet("Siswa: main solo (acak), live, atau misi.")
bullet("Orang tua: tidak masuk menu ini.")

h("2.2 Buat & terbitkan kuis", 2)
num("Buka tab Arena Belajar → Buat Kuis.")
num("Isi judul, soal (PG / benar-salah / isian / pasangkan), kunci.")
num("Simpan draf.")
num("Klik Terbitkan (ikuti petunjuk jari jika masih draf).")
num("Opsional: pilih skin template (Quiz, Pasangkan, Flashcard, Teka-teki, Susun kata, Ular tangga).")
add_img("arena-kuis.png")

h("2.3 Salin soal ke kelas lain", 2)
num("Buka experience.")
num("Bagian Salin soal ke kelas lain → centang kelas (mapel sama).")
num("Konfirmasi — experience baru dibuat di kelas tujuan.")

h("2.4 Siswa main solo", 2)
num("Siswa buka Arena → pilih kuis terbit → Mulai.")
num("Perhatikan chip Solo · soal acak (urutan soal & opsi beda tiap attempt).")
num("Jawab → Kumpulkan → lihat hasil.")

h("2.5 Live di kelas", 2)
num("Guru: Mulai Live → tunggu lobi → mulai soal.")
num("Maju soal sebagai host → Akhiri → podium/hasil.")
num("Transfer Nilai ke formatif/sumatif bila siap.")

h("2.6 Misi edukatif", 2)
num("Mode Misi → filter jenjang SD/SMP/SMA-SMK atau Tren 25–26.")
num("Guru tugaskan misi → siswa main → monitor → transfer nilai.")
num("Debrief singkat setelah main.")
add_img("arena-misi.png")

h("2.7 Video Arena", 2)
bullet("Ringkasan hub: arenabelajar.mp4")
bullet("Kuis/template/salin/solo: arenakuis.mp4")
bullet("Misi: arenamisi.mp4")

h("3. Asisten Guru", 1)
p("Lokasi: Akademik → Asisten Guru (bukan untuk siswa/orang tua).", bold=True)
add_img("asisten-ai.png")

h("3.1 Setup API key (sekali)", 2)
num("Buka Asisten Guru → modal Hubungkan API key Gemini.")
num("Buka Google AI Studio → Create API key → salin.")
num("Tempel di SIMS → Simpan API key.")
p("Jangan bagikan API key. Ganti/hapus lewat folder Nalar Guru & Kuota.", bold=True)

h("3.2 Folder Nalar Guru & Kuota (update)", 2)
bullet("Nalar (chat) dan status Generate Kuota digabung dalam satu folder collapsible.")
bullet("Pakai saran prompt atau ketik pertanyaan sendiri.")
bullet("Pantau sisa kuota LIVE sebelum generate banyak.")

h("3.3 Tab generator", 2)
bullet("Generator Soal — dari topik atau unggah materi.")
bullet("RPM Learning — rancangan pembelajaran.")
bullet("Perangkum Materi — ringkas bahan ajar.")
bullet("Draft Feedback — draf umpan balik siswa.")
bullet("History Generate — buka ulang hasil lama.")

h("3.4 Alur mengajar → Kirim ke Arena", 2)
num("Buat soal di Nalar / Generator.")
num("Periksa kunci & kesesuaian kurikulum.")
num("Kartu Alur mengajar → Kirim ke Arena → pilih ruang kelas.")
num("Form buat kuis Arena terbuka dengan soal terimpor → sunting bila perlu → Terbitkan.")
num("Opsional: Studio Presentasi / Canva Pendidikan.")

h("3.5 Video Asisten Guru", 2)
bullet("ai.mp4 — setup key, folder Nalar & Kuota, generator, kirim ke Arena.")

h("4. Checklist cepat sebelum kelas", 1)
bullet("API key Gemini tersimpan & kuota cukup.")
bullet("Soal sudah diperiksa manusia (bukan mentah dari AI).")
bullet("Experience Arena sudah Terbitkan.")
bullet("Kalau live: perangkat siswa online; kalau solo: jendela kuis terbuka.")
bullet("Setelah selesai: cek hasil → transfer nilai bila perlu.")

h("5. Cara membuka panduan di SIMS", 1)
num("Login sebagai guru.")
num("Menu Bantuan → Panduan Visual.")
num("Cari “Arena” atau “Asisten Guru”, putar video di kartu fitur.")

foot = doc.add_paragraph()
r = foot.add_run(
    "Dokumen ini melengkapi docs/PANDUAN_PENGGUNAAN_SIMS_APP.md dan resources/panduan/visual.html. "
    "Render ulang video: python docs/panduan-render/render_arena_ai_videos.py"
)
font(r, size=9, color=(120, 120, 120))

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
